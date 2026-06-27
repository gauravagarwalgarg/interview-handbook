#!/usr/bin/env python3
"""ATS Scanner  Matches resumes against JD keywords."""

import os
import re
import subprocess
from pathlib import Path
from dataclasses import dataclass, field
from docx import Document

RESUME_DIR = Path(__file__).parent.parent / "Resume"
JD_PATH = RESUME_DIR / "JD.md"

# === JD Keywords (extracted from JD.md) ===
# Weighted: core skills get higher weight, nice-to-haves get lower
JD_KEYWORDS = {
    # Core: C/C++ (weight 3)
    "c++": 3, "cpp": 3, "modern c++": 3, "c++14": 3, "c++17": 3, "c++20": 3,
    "stl": 3, "raii": 2, "smart pointer": 2, "template": 2,
    # Core: Build systems (weight 2)
    "cmake": 2, "ninja": 2, "make": 1,
    # Core: Embedded Linux (weight 3)
    "embedded": 3, "linux": 3, "embedded linux": 3, "rtos": 2,
    "device driver": 2, "kernel": 2,
    # Core: Yocto (weight 3)
    "yocto": 3, "bitbake": 3, "openembedded": 2, "meta-": 2, "recipe": 2,
    "layer": 1, "poky": 2, "petalinux": 2,
    # Core: Communication protocols (weight 2)
    "can": 2, "uart": 2, "spi": 2, "i2c": 2,
    "wi-fi": 1, "wifi": 1, "bluetooth": 1, "ble": 1,
    "ethernet": 1, "tcp": 1, "udp": 1,
    # Core: Networking/middleware (weight 2)
    "rest": 2, "http": 1, "rpc": 2, "grpc": 2, "mqtt": 2,
    "ipc": 2, "socket": 1, "protobuf": 2, "json": 1,
    # Core: CI/CD (weight 2)
    "gitlab": 2, "ci/cd": 2, "ci cd": 2, "pipeline": 2,
    "docker": 2, "container": 1, "jenkins": 1,
    # Core: Quality/Testing (weight 2)
    "sonarqube": 2, "clang-tidy": 2, "static analysis": 2,
    "unit test": 2, "gtest": 2, "pytest": 2, "test automation": 2,
    "code coverage": 1, "gcov": 1,
    # Core: Python (weight 2)
    "python": 2, "scripting": 1, "automation": 1,
    # Core: Architecture/Design (weight 2)
    "design pattern": 2, "architecture": 2, "oop": 1,
    "object oriented": 1,
    # Core: Agile (weight 1)
    "agile": 1, "scrum": 1, "sprint": 1, "kanban": 1,
    # Nice-to-have: AWS/Cloud (weight 1)
    "aws": 1, "cloud": 1, "s3": 1, "ec2": 1, "lambda": 1,
    # Nice-to-have: Web (weight 1)
    "javascript": 1, "web": 1, "react": 1, "node": 1, "html": 1,
    # Nice-to-have: Metrics (weight 1)
    "metrics": 1, "quality gate": 1,
    # Platform-specific (weight 2)
    "arm": 2, "cortex": 2, "zynq": 2, "xilinx": 2, "fpga": 2,
    # Version control (weight 1)
    "git": 1,
}

# Maximum possible score
MAX_SCORE = sum(JD_KEYWORDS.values())


def extract_text_pdf(path: Path) -> str:
    """Extract text from PDF using pdftotext."""
    try:
        result = subprocess.run(
            ["pdftotext", "-layout", str(path), "-"],
            capture_output=True, text=True, timeout=30
        )
        return result.stdout
    except Exception as e:
        return ""


def extract_text_docx(path: Path) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except Exception as e:
        print(f"  ⚠️  DOCX extraction error for {path.name}: {e}")
        return ""


def extract_text(path: Path) -> str:
    """Extract text based on file extension."""
    if path.suffix.lower() == ".pdf":
        return extract_text_pdf(path)
    elif path.suffix.lower() == ".docx":
        return extract_text_docx(path)
    return ""


def score_resume(text: str) -> tuple[float, dict, dict]:
    """Score resume text against JD keywords. Returns (pct, matched, missed)."""
    text_lower = text.lower()
    matched = {}
    missed = {}
    
    for keyword, weight in JD_KEYWORDS.items():
        # Check for keyword presence (word boundary for short keywords)
        if len(keyword) <= 3:
            pattern = r'\b' + re.escape(keyword) + r'\b'
            if re.search(pattern, text_lower):
                matched[keyword] = weight
            else:
                missed[keyword] = weight
        else:
            if keyword in text_lower:
                matched[keyword] = weight
            else:
                missed[keyword] = weight
    
    score = sum(matched.values())
    pct = (score / MAX_SCORE) * 100
    return pct, matched, missed


def get_candidate_name(filename: str) -> str:
    """Extract candidate name from filename."""
    name = filename
    # Remove zero-width characters
    name = re.sub(r'[\u200b\u200c\u200d\ufeff]', '', name)
    # Remove common suffixes
    for suffix in [".pdf", ".docx", "Zone.Identifier", "_Resume_", "_Resume", "Resume_", "Resume", "resume", "CV_", "cv_"]:
        name = name.replace(suffix, "")
    # Clean up
    name = name.replace("_", " ").replace("-", " ").strip()
    # Remove leading numbers/timestamps
    name = re.sub(r'^\d[\d\s]*\d\s*', '', name)
    # Remove trailing version/year info
    name = re.sub(r'\s*(v\d+|20\d{2}|\d+Yrs?|\d+\.\d+).*$', '', name, flags=re.IGNORECASE)
    return name.strip() or "(unknown)"


@dataclass
class CandidateResult:
    name: str
    filename: str
    score_pct: float
    matched_keywords: dict = field(default_factory=dict)
    missed_keywords: dict = field(default_factory=dict)
    experience_years: str = ""


def extract_experience(text: str) -> str:
    """Try to extract years of experience from resume text."""
    patterns = [
        r'(\d+\.?\d*)\s*\+?\s*years?\s*(of)?\s*(experience|exp)',
        r'experience.*?(\d+\.?\d*)\s*\+?\s*years?',
        r'(\d+\.?\d*)\s*years?\s*(of)?\s*(professional|industry|work)',
    ]
    for pattern in patterns:
        m = re.search(pattern, text.lower())
        if m:
            return m.group(1) + " years"
    return "N/A"


def main():
    results: list[CandidateResult] = []
    
    resume_files = [
        f for f in RESUME_DIR.iterdir()
        if f.suffix.lower() in (".pdf", ".docx") and "Zone.Identifier" not in f.name
    ]
    
    print(f"{'='*70}")
    print(f"  ATS SCAN REPORT  {len(resume_files)} Resumes vs JD")
    print(f"  Max Possible Score: {MAX_SCORE} points")
    print(f"  Threshold: 25%")
    print(f"{'='*70}\n")
    
    for resume_path in sorted(resume_files):
        text = extract_text(resume_path)
        if not text.strip():
            print(f"  ⚠️  Could not extract text: {resume_path.name}")
            continue
        
        pct, matched, missed = score_resume(text)
        exp = extract_experience(text)
        name = get_candidate_name(resume_path.name)
        
        results.append(CandidateResult(
            name=name,
            filename=resume_path.name,
            score_pct=pct,
            matched_keywords=matched,
            missed_keywords=missed,
            experience_years=exp,
        ))
    
    # Sort by score descending
    results.sort(key=lambda r: r.score_pct, reverse=True)
    
    # Print ranked results
    print(f"{'#':<4} {'Candidate':<35} {'Score':<8} {'Exp':<12} {'Status'}")
    print(f"{'-'*4} {'-'*35} {'-'*8} {'-'*12} {'-'*10}")
    
    for i, r in enumerate(results, 1):
        status = "✅ PASS" if r.score_pct >= 25 else "❌ BELOW"
        print(f"{i:<4} {r.name:<35} {r.score_pct:5.1f}%  {r.experience_years:<12} {status}")
    
    # Detailed report for passing candidates
    passing = [r for r in results if r.score_pct >= 25]
    
    print(f"\n{'='*70}")
    print(f"  DETAILED REPORT  {len(passing)} Candidates Above 25% Threshold")
    print(f"{'='*70}\n")
    
    for r in passing:
        print(f"┌─ {r.name} ({r.score_pct:.1f}%) ─ {r.experience_years}")
        print(f"│  File: {r.filename}")
        print(f"│")
        print(f"│  ✅ Matched Keywords ({len(r.matched_keywords)}):")
        # Group by weight
        high = {k: v for k, v in r.matched_keywords.items() if v >= 3}
        med = {k: v for k, v in r.matched_keywords.items() if v == 2}
        low = {k: v for k, v in r.matched_keywords.items() if v == 1}
        if high:
            print(f"│     CORE:  {', '.join(sorted(high.keys()))}")
        if med:
            print(f"│     IMPORTANT: {', '.join(sorted(med.keys()))}")
        if low:
            print(f"│     NICE-TO-HAVE: {', '.join(sorted(low.keys()))}")
        print(f"│")
        
        # Critical misses (weight >= 2)
        critical_miss = {k: v for k, v in r.missed_keywords.items() if v >= 2}
        if critical_miss:
            print(f"│  ❌ Critical Gaps (weight ≥ 2):")
            print(f"│     {', '.join(sorted(critical_miss.keys())[:15])}")
        print(f"└{'─'*68}\n")
    
    # Summary
    print(f"\n{'='*70}")
    print(f"  SUMMARY")
    print(f"{'='*70}")
    print(f"  Total resumes scanned: {len(results)}")
    print(f"  Above threshold (≥25%): {len(passing)}")
    print(f"  Below threshold (<25%): {len(results) - len(passing)}")
    if passing:
        print(f"  Top candidate: {passing[0].name} ({passing[0].score_pct:.1f}%)")
    print()


if __name__ == "__main__":
    main()
